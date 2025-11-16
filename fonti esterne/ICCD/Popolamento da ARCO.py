import re
import json
import argparse
from SPARQLWrapper import SPARQLWrapper, JSON
from docx import Document
from collections import Counter
from collections import defaultdict
from tqdm import tqdm
from Utils import ENDPOINT, generate_tags, get_city_from_args, Spinner

city_lowered, city_capitalized = get_city_from_args()


# Query principale per ottenere POI e URI di dating
# Attualmente la query è esplicitamente limitata al grafo GRAPH <https://w3id.org/arco/data> { ... }
# Questo significa che il motore SPARQL cerca solo nel dataset “ArCo”, cioè quello dell’Istituto Centrale per il Catalogo e la Documentazione (ICCD) — https://w3id.org/arco
sparql_query = f"""
PREFIX cd: <https://w3id.org/arco/ontology/construction-description/>
SELECT 
  ?s AS ?s_norm
  (SAMPLE(STR(?l)) AS ?label)
  (SAMPLE(?Categoria) AS ?Categoria)
  (SAMPLE(?Indirizzo) AS ?Indirizzo)
  (SAMPLE(COALESCE(?LatGeom, ?Lat)) AS ?Latitudine)
  (SAMPLE(COALESCE(?LongGeom, ?Long)) AS ?Longitudine)
  (SAMPLE(?Descrizione) AS ?Descrizione)
  (GROUP_CONCAT(DISTINCT STR(?Dating); separator=" | ") AS ?DatingURIs)
  (GROUP_CONCAT(DISTINCT STR(?constructionElement); separator=" | ") AS ?ConstructionElementURIs)
  (GROUP_CONCAT(DISTINCT STR(?covering); separator=" | ") AS ?CoveringURIs)
  (GROUP_CONCAT(DISTINCT STR(?design); separator=" | ") AS ?DesignURIs)
  (GROUP_CONCAT(DISTINCT STR(?layout); separator=" | ") AS ?LayoutURIs)
WHERE {{
  GRAPH <https://w3id.org/arco/data> {{
    ?s a arco:ImmovableCulturalProperty ;
       rdfs:label ?l ;
       a-loc:hasCulturalPropertyAddress ?address .
    ?address clvapit:hasCity ?city .
    ?city rdfs:label ?cityLabel .

    FILTER(LCASE(STR(?cityLabel)) = "{city_lowered}")

    OPTIONAL {{ ?address clvapit:fullAddress ?Indirizzo }}
    OPTIONAL {{ ?address geo:lat ?Lat }}
    OPTIONAL {{ ?address geo:long ?Long }}
    OPTIONAL {{ 
      ?s clvapit:hasGeometry ?geom .
      ?geom a-loc:hasCoordinates ?coords .
      OPTIONAL {{ ?coords a-loc:lat ?LatGeom }}
      OPTIONAL {{ ?coords a-loc:long ?LongGeom }} 
    }}
    OPTIONAL {{ ?s dc:description ?Descrizione }}
    OPTIONAL {{ ?s dc:type ?Categoria }}
    OPTIONAL {{ ?s a-cd:hasDating ?Dating }}

    # Architettura / Arte / Forma: manteniamo solo gli URI
    OPTIONAL {{ ?s cd:hasConstructionElement ?constructionElement }}
    OPTIONAL {{ ?s cd:hasCovering ?covering }}
    OPTIONAL {{ ?s cd:hasDesign ?design }}
    OPTIONAL {{ ?s cd:hasLayout ?layout }}

    BIND(REPLACE(STR(?s), ".*/", "") AS ?rawId)
    BIND(REPLACE(?rawId, "[^0-9]+$", "") AS ?idBase)
  }}
}}
GROUP BY ?idBase ?s
"""


categoria_map = {
    # Architetture religiose
    "chiesa": "Chiesa",
    "basilica": "Basilica / Duomo",
    "cattedrale": "Basilica / Duomo",
    "cappella": "Cappella / Oratorio",
    "oratorio": "Cappella / Oratorio",
    "monastero": "Monastero / Convento",
    "convento": "Monastero / Convento",
    "canonica": "Monastero / Convento",

    # Architettura civile e militare / Palazzo storico
    "palazzo": "Palazzo storico",
    "villa": "Palazzo storico",
    "castello": "Castello / Fortezza",
    "torre": "Torre / Campanile",
    "porta": "Porta / Mura urbane",
    "mura": "Porta / Mura urbane",
    "piazza": "Piazza / Spazio urbano",
    "ponte": "Ponte / Acquedotto",
    "acquedotto": "Ponte / Acquedotto",
    "convitto": "Palazzo storico",
    "casa privata": "Palazzo storico",
    "casa": "Palazzo storico",
    "edificio residenziale": "Palazzo storico",
    "rifugio": "Palazzo storico",
    "scuola": "Palazzo storico",
    "accademia": "Palazzo storico",
    "orfanotrofio": "Palazzo storico",
    "canonico": "Palazzo storico",
    "mercato": "Piazza / Spazio urbano",
    "architettura dello stato": "Palazzo storico",
    "edificio-residenziale": "Palazzo storico",
       
    # Musei e Collezioni
    "museo": "Museo",
    "galleria": "Mostra / Galleria",
    "mostra": "Mostra / Galleria",

    # Archeologia
    "area archeologica": "Area archeologica",
    "monumento antico": "Monumento antico",

    # Arte e Scultura
    "statua": "Statua / Monumento",
    "monumento": "Statua / Monumento",
    "fontana": "Fontana",
    "opera d’arte": "Opera d’arte",
    "dipinto": "Opera d’arte",
    "affresco": "Opera d’arte",
    "scultura": "Opera d’arte",
    "mosaico": "Opera d’arte",

    # Paesaggio e Natura
    "parco": "Parco / Giardino",
    "giardino": "Parco / Giardino",
    "belvedere": "Belvedere / Panorama",
    "panorama": "Belvedere / Panorama",
    "riserva naturale": "Riserva naturale",

    # Memoria e Identità
    "monumento ai caduti": "Luogo simbolico / Memoria",
    "lapide": "Luogo simbolico / Memoria",

    # Area urbana
    "centro storico": "Piazza / Spazio urbano",
    
    # Convitto / Casa privata ecc.
    "convittorio": "Palazzo storico",
    "casa-privata": "Palazzo storico",

    # Contemporaneo e urbano
    "installazione": "Opera contemporanea",
    "arredo urbano": "Elemento urbano (generico)",

    # Beni immateriali e territoriali
    "tradizione": "Tradizione / Rituale / Folklore",
    "rituale": "Tradizione / Rituale / Folklore",
    "folklore": "Tradizione / Rituale / Folklore",
    "percorso urbano": "Percorso urbano esperienziale",
    "itinerario culturale": "Itinerario culturale territoriale",

    # Servizi e spazi contemporanei
    "centro eventi": "Centro eventi / Teatro / Auditorium",
    "teatro": "Centro eventi / Teatro / Auditorium",
    "auditorium": "Centro eventi / Teatro / Auditorium",
    "spazio pubblico": "Spazio pubblico contemporaneo",
}

def run_sparql_query(endpoint, query):
    sparql = SPARQLWrapper(endpoint)
    sparql.setQuery(query)
    sparql.setReturnFormat(JSON)
    try:
        results = sparql.query().convert()
        return results["results"]["bindings"]
    except Exception as e:
        print("Errore SPARQL:", e)
        return []

def fetch_dating_details(dating_uri):
    """
    Restituisce solo label in italiano, time e description
    dai Dating URI, eliminando duplicati.
    """
    query = f"""
    SELECT ?label ?lang ?time ?timeLabel ?altTime ?descr
    WHERE {{
      GRAPH <https://w3id.org/arco/data> {{
        <{dating_uri}> <https://w3id.org/arco/ontology/context-description/hasDatingEvent> ?event .
        OPTIONAL {{ ?event <http://www.w3.org/2000/01/rdf-schema#label> ?label . }}
        OPTIONAL {{
          ?event <https://w3id.org/arco/ontology/context-description/specificTime> ?time .
          OPTIONAL {{ ?time <http://www.w3.org/2000/01/rdf-schema#label> ?timeLabel . }}
        }}
        OPTIONAL {{ ?event <http://www.w3.org/2006/time#atTime> ?altTime . }}
        OPTIONAL {{ ?event <https://w3id.org/arco/ontology/core/description> ?descr . }}
      }}
      BIND(LANG(?label) AS ?lang)
    }}
    """

    sparql = SPARQLWrapper(ENDPOINT)
    sparql.setQuery(query)
    sparql.setReturnFormat(JSON)

    try:
        res = sparql.query().convert()
        bindings = res.get("results", {}).get("bindings", [])

        triples = []
        seen = set()  # per eliminare duplicati basati solo su descr

        # Itera i binding in ordine inverso
        for r in bindings[::-1]:
            label = r.get("label", {}).get("value", "")
            lang = r.get("lang", {}).get("value", "")

            # Prendi solo label in italiano
            if lang != "it" or not label:
                continue

            # time: preferenza rdfs:label, altrimenti URI pulito
            if r.get("timeLabel", {}).get("value"):
                time = r["timeLabel"]["value"]
            elif r.get("time", {}).get("value"):
                time = r["time"]["value"].split("/")[-1].replace("-", " - ")
            else:
                time = ""

            altTime = r.get("altTime", {}).get("value", "")
            descr = r.get("descr", {}).get("value", "")

            # usa solo descr come chiave per i duplicati
            if descr in seen:
                continue
            seen.add(descr)

            triples.append((label, time or altTime, descr))

        return triples

    except Exception as e:
        print(f"Errore fetch_dating_details per {dating_uri}: {e}")
        return []

def fetch_construction_details(uri):
    """
    Recupera informazioni dettagliate da un URI Architettura/Arte/Forma.
    Costruisce testo leggibile dai campi principali: rdf:type, core:hasType, arco-dd:hasMaterial, arco-dd:hasShape (opzionali)
    """
    query = f"""
    PREFIX rdf: <http://www.w3.org/1999/02/22-rdf-syntax-ns#>
    PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>
    PREFIX arco-core: <https://w3id.org/arco/ontology/core/>
    PREFIX arco-dd: <https://w3id.org/arco/ontology/denotative-description/>
    PREFIX l0: <https://w3id.org/italia/onto/l0/>

    SELECT (CONCAT(
               STR(?elementTypeLabel),
               IF(BOUND(?typeLabelFinal), CONCAT(" ", STR(?typeLabelFinal)), ""),
               IF(BOUND(?materialLabelFinal), CONCAT(" ", STR(?materialLabelFinal)), ""),
               IF(BOUND(?shapeLabelFinal), CONCAT(" ", STR(?shapeLabelFinal)), "")
           ) AS ?fullLabel)
    WHERE {{
      <{uri}> rdf:type ?elementType .

      # label tipo principale
      OPTIONAL {{
        ?elementType rdfs:label ?elementTypeLabel .
        FILTER(lang(?elementTypeLabel)="it")
      }}
      OPTIONAL {{
        ?elementType l0:name ?elementTypeName .
      }}
      BIND(COALESCE(?elementTypeLabel, ?elementTypeName) AS ?elementTypeLabel)

      # tipo tecnico
      OPTIONAL {{ 
        <{uri}> arco-core:hasType ?type .
        OPTIONAL {{ ?type rdfs:label ?typeLabel FILTER(lang(?typeLabel)="it") }}
        OPTIONAL {{ ?type l0:name ?typeLabelName }}
        BIND(COALESCE(?typeLabel, ?typeLabelName) AS ?typeLabelFinal)
      }}

      # materiale
      OPTIONAL {{ 
        <{uri}> arco-dd:hasMaterial ?material .
        OPTIONAL {{ ?material rdfs:label ?materialLabel FILTER(lang(?materialLabel)="it") }}
        OPTIONAL {{ ?material l0:name ?materialLabelName }}
        BIND(COALESCE(?materialLabel, ?materialLabelName) AS ?materialLabelFinal)
      }}

      # forma
      OPTIONAL {{
        <{uri}> arco-dd:hasShape ?shape .
        OPTIONAL {{ ?shape rdfs:label ?shapeLabel FILTER(lang(?shapeLabel)="it") }}
        OPTIONAL {{ ?shape l0:name ?shapeLabelName }}
        BIND(COALESCE(?shapeLabel, ?shapeLabelName) AS ?shapeLabelFinal)
      }}
    }}
    LIMIT 1
    """

    sparql = SPARQLWrapper(ENDPOINT)
    sparql.setQuery(query)
    sparql.setReturnFormat(JSON)

    try:
        res = sparql.query().convert()
        bindings = res.get("results", {}).get("bindings", [])
        result_text = []

        for b in bindings:
            if "fullLabel" in b:
                result_text.append(b["fullLabel"]["value"])

        return result_text

    except Exception as e:
        print(f"Errore fetch_construction_details per {uri}: {e}")
        return []

def categorize_architecture_elements(elements):
    """
    Pulisce e raggruppa elementi di Architettura/Arte/Forma in categorie principali,
    restituendo testo pronto da inserire in un documento.
    """
    categories = defaultdict(list)

    for elem in elements:
        # 1. Rimuove simboli come '(?)'
        elem_clean = re.sub(r'\(\?\)', '', elem)
        # 2. Rimuove spazi multipli e strip
        elem_clean = re.sub(r'\s+', ' ', elem_clean).strip()
        
        # 3. Determina la categoria in base alla prima parola
        if elem_clean.lower().startswith("elemento orizzontale"):
            cat = "Elemento orizzontale"
            value = elem_clean[len("Elemento orizzontale "):].strip()
        elif elem_clean.lower().startswith("elemento verticale"):
            cat = "Elemento verticale"
            value = elem_clean[len("Elemento verticale "):].strip()
        elif elem_clean.lower().startswith("copertura"):
            cat = "Copertura"
            value = elem_clean[len("Copertura "):].strip()
        elif elem_clean.lower().startswith("pavimentazione"):
            cat = "Pavimentazione"
            value = elem_clean[len("Pavimentazione "):].strip()
        else:
            cat = "Altro"
            value = elem_clean
        
        # 4. Normalizza in minuscolo
        value = value.lower()
        categories[cat].append(value)
    
    # 5. Costruisce il testo finale con indentazione e elenco puntato
    output_lines = []
    for cat in ["Elemento orizzontale", "Elemento verticale", "Copertura", "Pavimentazione", "Altro"]:
        if cat in categories:
            output_lines.append(f"{cat}:")
            for v in categories[cat]:
                output_lines.append(f"  - {v}")
    
    return "\n".join(output_lines)

def clean_poi_name(nome_poi: str) -> str:
    """Rimuove la parte dopo il trattino nel Nome POI."""
    if '-' in nome_poi:
        nome_poi = nome_poi.split('-')[0].strip()     
    return nome_poi


def create_docx(data, filename=None):
    if filename is None:
        safe_city = city_capitalized.replace(" ", "_")
        filename = f"Dati_ARCO_{safe_city}.docx"
        
    doc = Document()
    doc.add_heading(f"Dati da ARCO – {city_capitalized}", level=0)

    for poi in tqdm(data, desc="Elaborazione POI", unit="POI"):
#    for poi in data:
        doc.add_heading(clean_poi_name(poi.get("label", {}).get("value", "POI sconosciuto")), level=1)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Nome Campo"
        hdr_cells[1].text = "Valore"

        # POI ID
        row = table.add_row().cells
        row[0].text = "POI ID"
        row[1].text = poi.get("s_norm", {}).get("value", "")

        # Nome POI
        row = table.add_row().cells
        row[0].text = "Nome POI"
        row[1].text = clean_poi_name(poi.get("label", {}).get("value", ""))

        # Categoria Persistente
        row = table.add_row().cells
        row[0].text = "Categoria Persistente"
        arco_categoria = poi.get("Categoria", {}).get("value", "")
        categoria_persistente = map_categoria(arco_categoria)
        row[1].text = categoria_persistente        

        # Storia e cronologia
        row = table.add_row().cells
        row[0].text = "Storia e cronologia"
        storia_text = ""
        descr_seen = set()
        dating_uris = poi.get("DatingURIs", {}).get("value", "").split(" | ")
        for uri in dating_uris:
            uri = uri.strip()
            if not uri:
                continue
            triples = fetch_dating_details(uri)
            for label, time, descr in triples:
                if any([label, time, descr]) and descr not in descr_seen:
                    descr_seen.add(descr)
                    storia_text += f"{label}, {time}\n{descr}\n\n"
        row[1].text = storia_text.strip()


        # Architettura / Arte / Forma
        row = table.add_row().cells
        row[0].text = "Architettura / Arte / Forma"

        arch_texts = []

        # unisci tutti i campi URI
        for key in ["ConstructionElementURIs", "CoveringURIs", "DesignURIs", "LayoutURIs"]:
            val = poi.get(key, {}).get("value")
            if not val:
                continue
            if isinstance(val, str):
                arch_texts += [u.strip() for u in val.split(" | ") if u.strip()]
            elif isinstance(val, list):
                arch_texts += [u.strip() for u in val if u.strip()]

        # recupera i dettagli da ogni URI
        arch_details = []
        for uri in arch_texts:
            result = fetch_construction_details(uri)
            if not result:
                print(f"Attenzione: nessun dato trovato per {uri}")
            else:
                arch_details.extend(result)

        # inserisci nel documento
        architettura_text = categorize_architecture_elements(arch_details) if arch_details else "-"
        row[1].text = architettura_text
        
        # --- RIGA TAG ---
        row = table.add_row().cells
        row[0].text = "Tag"

        # integra architettura_text nella generazione dei tag
        arch_text_for_tags = ""
        if architettura_text and architettura_text != "-":
            # rimuovi le intestazioni generiche
            arch_text_cleaned = re.sub(
                r"^(?:\s*-?\s*)?(Elemento orizzontale|Elemento verticale|Copertura|Pavimentazione|Altro):\s*", 
                "", 
                architettura_text, 
                flags=re.MULTILINE
            )
            # sostituisci punteggiatura e newline con spazi
            arch_text_cleaned = re.sub(r"[:\n\-]", " ", arch_text_cleaned)
            arch_text_cleaned = re.sub(r"\s+", " ", arch_text_cleaned).strip()
            arch_text_for_tags = arch_text_cleaned

        nome_poi = clean_poi_name(poi.get("label", {}).get("value", ""))
        descrizione = poi.get("Descrizione", {}).get("value", "")
        # genera tag escludendo la categoria persistente
        #tags_list = generate_tags(nome_poi, descrizione, storia_text.strip(), arch_text_for_tags, exclude_tags=[categoria_persistente])
        tags_list = generate_tags("", descrizione, storia_text.strip(), arch_text_for_tags, exclude_tags=[categoria_persistente])
        row[1].text = ", ".join(tags_list)        
        # print(f"tags: {row[1].text}")
        
        # Indirizzo
        row = table.add_row().cells
        row[0].text = "Indirizzo"
        row[1].text = poi.get("Indirizzo", {}).get("value", "")

        # Latitudine / Longitudine
        row = table.add_row().cells
        row[0].text = "Latitudine"
        row[1].text = poi.get("Latitudine", {}).get("value", "")
        row = table.add_row().cells
        row[0].text = "Longitudine"
        row[1].text = poi.get("Longitudine", {}).get("value", "")

    try:
        doc.save(filename)
        print(f"File salvato correttamente: {filename}")
    except PermissionError:
        print(f"Errore: impossibile salvare '{filename}'. Probabilmente il file è aperto in un'altra applicazione.")
    except Exception as e:
        print(f"Errore imprevisto durante il salvataggio: {e}")


def map_categoria(arco_categoria):
    """
    Mappa un valore qualsiasi di Categoria ARCO alla Categoria Persistente ICCD.
    Usa keyword matching: se una keyword è presente nel testo, restituisce la categoria ICCD corrispondente.
    Se non trova corrispondenza, restituisce il valore originale.
    """
    arco_categoria_lower = arco_categoria.lower()

    for keyword, categoria_iccd in categoria_map.items():
        if re.search(rf'\b{re.escape(keyword)}\b', arco_categoria_lower):
            return categoria_iccd

    print(f"fallback categoria: {arco_categoria}")
    # fallback: restituisci la Categoria originale della query
    return arco_categoria


if __name__ == "__main__":
    with Spinner("Esecuzione query SPARQL..."):
        results = run_sparql_query(ENDPOINT, sparql_query)
    
    create_docx(results)