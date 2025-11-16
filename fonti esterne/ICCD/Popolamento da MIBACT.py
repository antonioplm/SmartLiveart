import argparse
import requests
from SPARQLWrapper import SPARQLWrapper, JSON
from docx import Document
from tqdm import tqdm
from Utils import ENDPOINT, generate_tags, get_city_from_args, Spinner

city_lowered, city_capitalized = get_city_from_args()

def get_bounding_box(city_name, country="Italy"):
    url = "https://nominatim.openstreetmap.org/search"
    params = {"q": f"{city_name}, {country}", "format": "json", "limit": 1}
    response = requests.get(url, params=params, headers={"User-Agent": "Python-SPARQL"})
    response.raise_for_status()
    data = response.json()
    if not data:
        raise ValueError(f"Città non trovata: {city_name}")
    bbox = data[0]["boundingbox"]
    minLat, maxLat = float(bbox[0]), float(bbox[1])
    minLon, maxLon = float(bbox[2]), float(bbox[3])
    return minLat, maxLat, minLon, maxLon

def build_mibact_query(city_name):
    minLat, maxLat, minLon, maxLon = get_bounding_box(city_name)
    query = f"""
    SELECT ?s ?Nome_Istituzionale ?Descrizione ?Indirizzo
           ?Latitudine ?Longitudine
           (GROUP_CONCAT(DISTINCT ?Prenotazioni; separator=" | ") AS ?PrenotazioniAggregata)
           (GROUP_CONCAT(DISTINCT ?Orari_di_apertura; separator=" | ") AS ?OrariAggregati)
           (GROUP_CONCAT(DISTINCT ?Biglietti; separator=" | ") AS ?BigliettiAggregati)
           (GROUP_CONCAT(DISTINCT ?Servizi; separator=" | ") AS ?ServiziAggregati)
           (GROUP_CONCAT(DISTINCT ?Telefono; separator=" | ") AS ?TelefoniAggregati)
           (GROUP_CONCAT(DISTINCT ?Email; separator=" | ") AS ?EmailAggregati)
           (GROUP_CONCAT(DISTINCT ?WebSite; separator=" | ") AS ?WebSiteAggregati)
    WHERE {{
      GRAPH <http://dati.beniculturali.it/mibact/luoghi> {{
        ?s rdf:type cis:CulturalInstituteOrSite ;
           cis:institutionalCISName ?Nome_Istituzionale .

        OPTIONAL {{ ?s l0:description ?Descrizione }}
        OPTIONAL {{ ?s cis:hasSite ?site . ?site cis:siteAddress ?address . ?address clvapit:fullAddress ?Indirizzo }}
        OPTIONAL {{ ?s geo:lat ?Latitudine }}
        OPTIONAL {{ ?s geo:long ?Longitudine }}
        OPTIONAL {{ ?s accessCondition:hasAccessCondition ?access . ?access rdf:type accessCondition:Booking . ?access rdfs:label ?Prenotazioni }}
        OPTIONAL {{ ?s accessCondition:hasAccessCondition ?accessHours . ?accessHours rdf:type accessCondition:OpeningHoursSpecification . ?accessHours l0:description ?Orari_di_apertura }}
        OPTIONAL {{ ?s cis:providesService [l0:name ?Servizi] }}
        OPTIONAL {{ ?s potapit:hasTicket ?ticket . ?ticket potapit:hasPriceSpecification ?priceSpec . ?priceSpec potapit:hasCurrencyValue ?Biglietti }}
        OPTIONAL {{ ?s smapit:hasOnlineContactPoint ?contactPoint . 
                   OPTIONAL {{ ?contactPoint smapit:hasTelephone [smapit:telephoneNumber ?Telefono] }}
                   OPTIONAL {{ ?contactPoint smapit:hasEmail [smapit:emailAddress ?Email] }}
                   OPTIONAL {{ ?contactPoint smapit:hasWebSite [smapit:URL ?WebSite] }} }}

        FILTER(
            (BOUND(?Latitudine) && BOUND(?Longitudine) &&
             ?Latitudine >= {minLat} && ?Latitudine <= {maxLat} &&
             ?Longitudine >= {minLon} && ?Longitudine <= {maxLon})
            ||
            (BOUND(?Indirizzo) && REGEX(?Indirizzo, " - {city_name}$", "i"))
        )
      }}
    }}
    GROUP BY ?s ?Nome_Istituzionale ?Descrizione ?Indirizzo ?Latitudine ?Longitudine
    """
    return query

def build_query():
    city = city_lowered.replace('"', '\\"')  # escape degli eventuali doppi apici

    query = """
SELECT ?s ?Nome_Istituzionale ?Descrizione ?Indirizzo
       ?Latitudine ?Longitudine
       (GROUP_CONCAT(DISTINCT ?Prenotazioni; separator=" | ") AS ?PrenotazioniAggregata)
       (GROUP_CONCAT(DISTINCT ?Orari_di_apertura; separator=" | ") AS ?OrariAggregati)
       (GROUP_CONCAT(DISTINCT ?Biglietti; separator=" | ") AS ?BigliettiAggregati)
       (GROUP_CONCAT(DISTINCT ?Servizi; separator=" | ") AS ?ServiziAggregati)
       (GROUP_CONCAT(DISTINCT ?Telefono; separator=" | ") AS ?TelefoniAggregati)
       (GROUP_CONCAT(DISTINCT ?Email; separator=" | ") AS ?EmailAggregati)
       (GROUP_CONCAT(DISTINCT ?WebSite; separator=" | ") AS ?WebSiteAggregati)
WHERE {{
  GRAPH <http://dati.beniculturali.it/mibact/luoghi> {{
    ?s rdf:type cis:CulturalInstituteOrSite ;
       cis:institutionalCISName ?Nome_Istituzionale .
    
    OPTIONAL {{ ?s l0:description ?Descrizione }}
    
    OPTIONAL {{ 
      ?s cis:hasSite [cis:siteAddress ?address] .
      OPTIONAL {{ ?address clvapit:fullAddress ?Indirizzo }}
    }}

    OPTIONAL {{ ?s geo:lat ?Latitudine }}
    OPTIONAL {{ ?s geo:long ?Longitudine }}

    OPTIONAL {{ ?s accessCondition:hasAccessCondition [rdf:type accessCondition:Booking ; rdfs:label ?Prenotazioni] }}
    OPTIONAL {{ ?s accessCondition:hasAccessCondition [rdf:type accessCondition:OpeningHoursSpecification ; l0:description ?Orari_di_apertura] }}
    OPTIONAL {{ ?s cis:providesService [l0:name ?Servizi] }}

    OPTIONAL {{
      ?s potapit:hasTicket ?ticket .
      ?offer potapit:includes ?ticket ;
             potapit:hasPriceSpecification [potapit:hasCurrencyValue ?Biglietti]
    }}

    OPTIONAL {{ 
      ?s smapit:hasOnlineContactPoint ?contactPoint .
      OPTIONAL {{ ?contactPoint smapit:hasTelephone [smapit:telephoneNumber ?Telefono] }}
      OPTIONAL {{ ?contactPoint smapit:hasEmail [smapit:emailAddress ?Email] }}
      OPTIONAL {{ ?contactPoint smapit:hasWebSite [smapit:URL ?WebSite] }} 
    }}

    FILTER(BOUND(?Indirizzo) && CONTAINS(LCASE(?Indirizzo), "{city}"))
  }}
}}
GROUP BY ?s ?Nome_Istituzionale ?Descrizione ?Indirizzo ?Latitudine ?Longitudine
"""
    return query.format(city=city)


# ---------------------------
# SPARQL QUERY EXECUTION
# ---------------------------
def run_sparql(query):
    sparql = SPARQLWrapper(ENDPOINT)
    sparql.setQuery(query)
    sparql.setReturnFormat(JSON)

    try:
        res = sparql.query().convert()
        return res["results"]["bindings"]
    except Exception as e:
        print("Errore SPARQL:", e)
        return []

def remove_duplicates(data):
    cleaned = {}
    for row in data:
        uri = row['s']['value']
        if uri not in cleaned:
            cleaned[uri] = row.copy()
        else:
            # Aggiorna campi aggregabili concatenando nuovi valori
            for key in ['PrenotazioniAggregata','OrariAggregati','BigliettiAggregati','ServiziAggregati','TelefoniAggregati','EmailAggregati','WebSiteAggregati']:
                val = row.get(key, {}).get('value', '')
                if val:
                    existing = cleaned[uri].get(key, {}).get('value', '')
                    # evita duplicati nello stesso campo
                    all_vals = set(existing.split(' | ')) if existing else set()
                    all_vals.update(val.split(' | '))
                    cleaned[uri][key] = {'type': 'literal', 'value': ' | '.join(sorted(all_vals))}
    return list(cleaned.values())

# ---------------------------
# DOCX CREATION
# ---------------------------
def create_doc(data, filename=None):

    safe_city = city_capitalized.replace(" ", "_")

    if not filename:
        filename = f"Dati_MiBACT_{safe_city}.docx"

    doc = Document()
    doc.add_heading(f"Dati MiBACT – {city_capitalized}", level=0)

    for row in tqdm(data, desc="Elaborazione POI", unit="POI"):
#    for row in data:

        nome = row.get("Nome_Istituzionale", {}).get("value", "")
        doc.add_heading(nome, level=1)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Campo"
        hdr[1].text = "Valore"

        def add(label, key):
            r = table.add_row().cells
            r[0].text = label
            r[1].text = row.get(key, {}).get("value", "")

        add("URI", "s")
        add("Nome istituzionale", "Nome_Istituzionale")
        add("Descrizione", "Descrizione")
        
        # --- AGGIUNTA RIGA TAG ---
        r = table.add_row().cells
        r[0].text = "Tag"

        # genera tag per il POI MiBACT usando nome e descrizione
        nome_poi = row.get("Nome_Istituzionale", {}).get("value", "")
        descrizione = row.get("Descrizione", {}).get("value", "")
        tags_list = generate_tags(
            nome_poi="",
            descrizione=descrizione,
            storia="",
            arch_text="",
            city_lowered=city_lowered
        )
        r[1].text = ", ".join(tags_list)
        
        add("Indirizzo", "Indirizzo")
        add("Latitudine", "Latitudine")
        add("Longitudine", "Longitudine")
        add("Prenotazioni", "PrenotazioniAggregata")
        add("Orari di apertura", "OrariAggregati")
        add("Biglietti", "BigliettiAggregati")
        add("Servizi", "ServiziAggregati")
        add("Telefono", "TelefoniAggregati")
        add("Email", "EmailAggregati")
        add("Sito web", "WebSiteAggregati")


    doc.save(filename)
    print(f"\nDocumento generato: {filename}")


# ---------------------------
# MAIN
# ---------------------------
if __name__ == "__main__":
    with Spinner("Esecuzione query SPARQL..."):
        query = build_mibact_query(city_lowered)
        results = run_sparql(query)
        results = remove_duplicates(results)

    create_doc(results)
